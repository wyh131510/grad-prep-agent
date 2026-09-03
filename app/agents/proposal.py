# -*- coding: utf-8 -*-
"""开题报告：模板解析、分块生成、手动修改、导出（md/docx）。"""
from __future__ import annotations

import io
import re

from ..llm.client import LLMError, llm
from ..llm.prompts import proposal_section_prompt
from ..schemas import ProposalSection, Task, TemplateInfo
from ..utils import clean_text, start_progress_ticker, truncate

DEFAULT_SECTIONS: list[tuple[str, str]] = [
    ("background", "课题背景与研究意义"),
    ("literature_review", "国内外研究现状"),
    ("objectives", "研究内容与目标"),
    ("methodology", "研究方案与技术路线"),
    ("feasibility", "可行性分析"),
    ("schedule", "进度安排"),
    ("references", "参考文献"),
]

# 模板章节关键词 → 分块 key（用于从学校模板自动检测结构）
SECTION_ALIASES = {
    "背景": "background", "意义": "background", "目的": "background",
    "现状": "literature_review", "综述": "literature_review", "国内外": "literature_review", "文献": "literature_review",
    "内容": "objectives", "目标": "objectives", "任务": "objectives",
    "方案": "methodology", "路线": "methodology", "方法": "methodology", "技术": "methodology",
    "可行性": "feasibility",
    "进度": "schedule", "安排": "schedule", "计划": "schedule",
    "参考": "references",
}


def _extract_text(filename: str, data: bytes) -> str:
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    if ext == "docx":
        import docx

        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if ext == "pdf":
        import pymupdf as fitz

        with fitz.open(stream=data, filetype="pdf") as doc:
            return "\n".join(p.get_text() for p in doc)
    return data.decode("utf-8", errors="replace")


def detect_sections(text: str) -> list[str]:
    keys: list[str] = []
    for line in text.splitlines():
        s = line.strip()
        if not s or len(s) > 40:
            continue
        # 标题形式：一、xxx （一）xxx 1. xxx 第一章 xxx
        if re.match(r"^[（(]?[一二三四五六七八九十\d]+[）)、.．、]", s) or "章" in s[:4]:
            for kw, key in SECTION_ALIASES.items():
                if kw in s and key not in keys:
                    keys.append(key)
    return keys if len(keys) >= 3 else [k for k, _ in DEFAULT_SECTIONS]


def parse_template_file(filename: str, data: bytes) -> TemplateInfo:
    text = _extract_text(filename, data)
    text = clean_text(text)
    if not text.strip():
        raise ValueError("模板文件为空或无法解析（支持 docx/pdf/md/txt）")
    return TemplateInfo(filename=filename, content_md=text[:20000], sections=detect_sections(text))


def get_sections(task_id: str) -> list[ProposalSection]:
    from ..store import repo

    tmpl = repo.get_template(task_id)
    keys = tmpl.sections if tmpl else [k for k, _ in DEFAULT_SECTIONS]
    repo.ensure_sections(task_id, DEFAULT_SECTIONS)
    by_key = {s.key: s for s in repo.get_sections(task_id)}
    return [by_key[k] for k in keys if k in by_key]


def _materials(task: Task) -> str:
    """收藏文献的结构化材料（开题生成的事实基础）。"""
    from ..store import repo

    _, papers = repo.list_papers(task.id, collected=True, sort="year", order="desc", limit=200)
    if not papers:
        return ""
    lines = []
    for i, p in enumerate(papers, 1):
        lines.append(f"[{i}] {p.title}（{p.year or '年份不详'}；{p.venue or p.source}）")
        zh = p.title_zh or ""
        abs_text = p.abstract_zh or p.abstract or ""
        if zh:
            lines.append(f"    题译：{zh}")
        lines.append(f"    摘要：{truncate(abs_text, 280)}")
        if p.summary and p.summary.relevance_to_topic:
            lines.append(f"    与选题关联：{truncate(p.summary.relevance_to_topic, 200)}")
    return "\n".join(lines)[:14000]


def _template_hint(task_id: str, key: str) -> str:
    from ..store import repo

    tmpl = repo.get_template(task_id)
    if not tmpl:
        return ""
    for kw, k in SECTION_ALIASES.items():
        if k == key and kw in tmpl.content_md:
            idx = tmpl.content_md.find(kw)
            return clean_text(tmpl.content_md[max(0, idx - 100): idx + 500])
    return truncate(clean_text(tmpl.content_md), 1500)


def generate_section(task: Task, key: str, instruction: str = "", job=None) -> str:
    from ..store import repo

    by_key = {s.key: s for s in repo.get_sections(task.id)}
    if key not in by_key:
        raise ValueError(f"未知分块：{key}")
    section = by_key[key]
    existing = {k: v.content for k, v in by_key.items() if k != key and v.content.strip()}
    prompt = proposal_section_prompt(
        key, section.title, task.topic, task.major, task.requirements,
        _template_hint(task.id, key), _materials(task), existing, instruction,
    )
    if job:
        job.update(message=f"正在生成「{section.title}」（调用大模型，可能需 30~90 秒）…")
    stop = start_progress_ticker(job, f"正在生成「{section.title}」")
    try:
        content = llm.chat(
            role="proposal",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=4096,
        ).strip()
    except LLMError as exc:
        raise RuntimeError(f"生成「{section.title}」失败：{exc}") from exc
    finally:
        stop()
    if not content:
        raise RuntimeError("生成内容为空")
    repo.put_section(task.id, key, section.title, content, "draft")
    if job:
        job.update(0.98, f"「{section.title}」生成完成。")
    return content


def generate_all(task: Task, job=None) -> dict:
    sections = get_sections(task.id)
    todo = [s for s in sections if not s.content.strip()]
    generated: list[str] = []
    failed: list[dict] = []
    for i, s in enumerate(todo):
        if job:
            job.update((i + 0.3) / max(1, len(todo)), f"正在生成「{s.title}」（{i + 1}/{len(todo)}）…")
        try:
            generate_section(task, s.key, "", job)
            generated.append(s.key)
        except Exception as exc:  # noqa: BLE001
            failed.append({"key": s.key, "error": str(exc)})
            if job:
                job.update(message=f"「{s.title}」生成失败：{exc}")
    return {"generated": generated, "failed": failed}


def assemble_text(task: Task) -> str:
    sections = get_sections(task.id)
    parts = [f"# {task.topic} 开题报告"]
    for s in sections:
        if s.content.strip():
            parts.append(s.content.strip())
    return "\n\n".join(parts)


def export_markdown(task: Task) -> str:
    return assemble_text(task) + "\n"


def export_docx(task: Task) -> bytes:
    import docx as docx_lib
    from docx.oxml.ns import qn
    from docx.shared import Pt

    FONT = "Microsoft YaHei"  # 微软雅黑：Windows 自带，中文字形完整（避免 Word 缺字形显示“口”）

    doc = docx_lib.Document()
    # 1) Normal 样式默认字体（中英文都用统一字体）
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    if normal._element.rPr is not None and normal._element.rPr.rFonts is not None:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    def _style_run(run, *, size=None, bold=None):
        run.font.name = FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), FONT)
        rfonts.set(qn("w:hAnsi"), FONT)
        rfonts.set(qn("w:eastAsia"), FONT)
        if size:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        return run

    # 2) 标题
    title_para = doc.add_heading(task.topic, level=0)
    for run in title_para.runs:
        _style_run(run, size=18, bold=True)

    # 3) 各分块
    for s in get_sections(task.id):
        if not s.content.strip():
            continue
        heading = doc.add_heading(s.title, level=1)
        for run in heading.runs:
            _style_run(run, size=14, bold=True)
        for para in (s.content.strip()).split("\n"):
            p = para.strip()
            if not p or p.startswith("##"):
                continue
            if p.startswith(("#", "-", "*")):
                para_obj = doc.add_paragraph(p.lstrip("#-* ").strip(), style="List Bullet")
            else:
                para_obj = doc.add_paragraph(p)
            for run in para_obj.runs:
                _style_run(run, size=11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
